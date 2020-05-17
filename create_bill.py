def info(phoneNum,IpNum):
    import int_billing,phone_billing
    int_invoice = int_billing.invoices(IpNum)
    a,b = phone_billing.invoices(phoneNum)
    phone_invoice = a+b
    ssum = int_invoice + phone_invoice
    tax = ssum * 0.2 # tax = 20% # это налог равно 20%
    tsum = ssum + tax
    int_invoice = '%.2f' % int_invoice
    phone_invoice = '%.2f' % phone_invoice
    ssum = '%.2f' % ssum
    tax = '%.2f' % tax
    tsum = '%.2f' % tsum
    # Добавление данных в шаблоне
    data = {}
    data['Банк получателя'] = 'Банк ВТБ 24'
    data['Получатель'] = 'Ле Тиен Зунг'
    data['ИНН'] = 'ИНН 123456890'
    data['КПП'] = 'КПП 686878912'
    data['БИК'] = '697895'
    data['1Сч. №'] = '12345678901234567890'
    data['2Сч. №'] = '24681357912345687690'
    data['3Сч. №'] = '68'
    data['день'] = '16 '
    data['месяц'] = 'Мая '
    data['год'] = '20'
    data['Поставщик'] = 'Банк ВТБ 24 ,ИНН 323456890,КПП 986878912,897895,Санкт-Петербург г., Вяземский переулок ,дом 5/7, тел.79522836868:'
    data['Заказчик'] = r'ООО "CS",ИНН 9714037378,КПП 787550001,159361,Санкт-Петербург г.,Васиевостров ул.,дом № 58'
    data['Основание'] = '№ 697896 от 13.03.2020'
    data['№'] = ['1','2']
    data['товары'] = ['Тарификация услуг "Телефония"','Тарификация услуг "Интернет"']
    data['количество'] = ['','']
    data['ед.'] = ['','']
    data['цена'] = [phone_invoice,int_invoice]
    data['сумма'] = [phone_invoice,int_invoice]
    data['Итого'] = ssum
    data['НДС'] = tax
    data['Всего к оплату'] = tsum
    data['Всего наименований'] = 'Всего наименований 2, на сумму '+ tsum +' руб.'
    data['Сумма прописью'] = ''
    return data
def pdf(phoneNum,IpNum):
    data = info(phoneNum,IpNum)
    import comtypes.client
    wdFormatPDF = 17

#Это мой локаль файл , в котором я сохранил 
    in_file = r'E:\2 курс\kypc 3\Mobile\lab3_beo\ITMO---mobile-devices\lab3\example.docx'
    out_file = r'E:\2 курс\kypc 3\Mobile\lab3_beo\ITMO---mobile-devices\lab3\bill.docx'
    pdf_file = r'E:\2 курс\kypc 3\Mobile\lab3_beo\ITMO---mobile-devices\lab3\bill.pdf'

    from docx import Document

    doc = Document(in_file)
    #{-----------------table 1----------------}
    table = doc.tables[0]
    table.cell(0,0).paragraphs[0].text = data['Банк получателя']
    table.cell(0,3).paragraphs[0].text = data['БИК']
    table.cell(2,3).paragraphs[0].text = data['1Сч. №']
    table.cell(3,0).paragraphs[0].text = data['ИНН']
    table.cell(3,1).paragraphs[0].text = data['КПП']
    table.cell(3,3).paragraphs[0].text = data['2Сч. №']
    table.cell(4,0).paragraphs[0].text = data['Получатель']
    #{---------------table 2--------------------}
    table = doc.tables[1]
    inline = table.cell(0,0).paragraphs[0].runs
    text = inline[1].text.replace('___',data['3Сч. №'])
    inline[1].text = text
    text = inline[3].text.replace('___',data['день'])
    inline[3].text = text
    text = inline[5].text.replace('_______',data['месяц'])
    inline[5].text = text
    text = inline[7].text.replace('___',data['год'])
    inline[7].text = text
    #{---------------table 3--------------------}
    table = doc.tables[2]
    table.cell(0,1).paragraphs[0].text = data['Поставщик']
    table.cell(1,1).paragraphs[0].text = data['Заказчик']
    table.cell(2,1).paragraphs[0].text = data['Основание']
    #{---------------table 4--------------------}
    table = doc.tables[3]
    for i in range(len(data['№'])):
        if i>0: table.add_row()
        table.cell(i+1,0).text = data['№'][i]
        table.cell(i+1,1).text = data['товары'][i]
        table.cell(i+1,2).text = data['количество'][i]
        table.cell(i+1,3).text = data['ед.'][i]
        table.cell(i+1,4).text = data['цена'][i]
        table.cell(i+1,5).text = data['сумма'][i]
    #{---------------bang 5--------------------}
    table = doc.tables[4]
    table.cell(0,1).paragraphs[0].text = data['Итого']
    table.cell(1,1).paragraphs[0].text = data['НДС']
    table.cell(2,1).paragraphs[0].text = data['Всего к оплату']
    table.cell(3,0).paragraphs[1].text = data['Всего наименований']
    table.cell(4,0).paragraphs[0].text = data['Сумма прописью']

    #{----------------save doc file--------------}
    doc.save(out_file)
    #{----------------create pdf file--------------}
    word = comtypes.client.CreateObject('Word.Application')
    doc = word.Documents.Open(out_file)
    doc.SaveAs(pdf_file, FileFormat=wdFormatPDF)
    doc.Close()
    word.Quit()

# вариант мой номер : 911926375 и с ip 192.168.250.41
pdf('911926375','192.168.250.41')